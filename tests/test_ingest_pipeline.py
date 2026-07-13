"""Test fixtures for the ingestion pipeline.

Creates 5 test files in a temp directory:
  1. text_guide.pdf      — text-based PDF with Chinese + English content
  2. scanned_doc.pdf      — scanned/image-only PDF (rendered as images)
  3. report.docx          — DOCX with paragraphs + a table
  4. notes.txt            — plain text file (UTF-8, Chinese + English)
  5. readme.md            — Markdown file with headings + code blocks

Usage:
  cd ~/rag-kit
  python tests/test_ingest_pipeline.py
"""

from __future__ import annotations

import os
import sys
import shutil
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from rag_kit.ingest import ingest_file, ingest_folder, SUPPORTED_EXTENSIONS


def _create_text_pdf(path: Path) -> None:
    """Create a text-based PDF with Chinese + English content."""
    import fitz

    doc = fitz.open()

    # Page 1: English content
    page1 = doc.new_page()
    page1.insert_text(
        (72, 72),
        "System Configuration Guide\n\n"
        "This document describes how to configure the testing system.\n\n"
        "1. Hardware Setup\n"
        "Connect the memory test module to the PCIe slot.\n"
        "Ensure stable power supply (minimum 750W recommended).\n\n"
        "2. Software Installation\n"
        "Install drivers before connecting the device.\n"
        "Supports Windows 10/11 and Linux operating systems.\n"
        "Restart the computer after installation to apply changes.",
        fontsize=11,
    )

    # Page 2: Chinese content
    page2 = doc.new_page()
    page2.insert_text(
        (72, 72),
        "系统配置指南\n\n"
        "本文档介绍如何配置测试系统。\n\n"
        "1. 硬件配置\n"
        "将内存测试模块连接到主板的PCIe插槽。\n"
        "确保电源供应稳定，建议至少750W。\n\n"
        "2. 软件安装\n"
        "安装驱动程序前请关闭所有测试程序。\n"
        "支持Windows 10/11和Linux操作系统。",
        fontsize=11,
        fontname="china-s",  # built-in CJK font in PyMuPDF
    )

    doc.save(str(path))
    doc.close()


def _create_scanned_pdf(path: Path) -> None:
    """Create a scanned/image-only PDF (no extractable text)."""
    import fitz
    from PIL import Image, ImageDraw, ImageFont

    # Create an image with rendered text (not extractable by PyMuPDF).
    img = Image.new("RGB", (595, 842), "white")
    draw = ImageDraw.Draw(img)

    # Try to use a default font.
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except (OSError, IOError):
        font = ImageFont.load_default()

    draw.text((72, 72), "Scanned Document", fill="black", font=font)
    draw.text((72, 120), "This is a scanned page.", fill="black", font=font)
    draw.text((72, 160), "OCR should extract this text.", fill="black", font=font)

    # Convert image to PDF page.
    png_bytes = _img_to_bytes(img)
    doc = fitz.open()
    page = doc.new_page(width=img.width, height=img.height)
    page.insert_image(rect=page.rect, stream=png_bytes)
    doc.save(str(path))
    doc.close()


def _img_to_bytes(img) -> bytes:
    """Convert a PIL image to PNG bytes."""
    import io
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _create_docx(path: Path) -> None:
    """Create a DOCX with paragraphs and a table."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("Technical Report", level=1)
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "This report covers the memory testing methodology "
        "used in our laboratory. The system supports both "
        "DDR4 and DDR5 memory modules."
    )
    doc.add_heading("Test Results", level=2)

    # Add a table.
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Module", "Speed (MT/s)", "Status"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    data = [
        ["DDR4-3200", "3200", "PASS"],
        ["DDR4-3600", "3600", "PASS"],
        ["DDR5-4800", "4800", "FAIL"],
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph("Conclusion: The system is functioning correctly.")
    doc.save(str(path))


def _create_txt(path: Path) -> None:
    """Create a plain text file with Chinese + English content."""
    content = (
        "故障排除手册\n"
        "常见问题及解决方法\n\n"
        "问题1：测试结果不准确\n"
        "可能原因：系统未校准。请按照配置指南进行系统校准。\n\n"
        "Issue 2: Device not recognized\n"
        "Possible cause: Drivers not properly installed.\n"
        "Please reinstall drivers and restart the system.\n"
    )
    path.write_text(content, encoding="utf-8")


def _create_md(path: Path) -> None:
    """Create a Markdown file with headings and code blocks."""
    content = (
        "# README\n\n"
        "## Overview\n\n"
        "This is a test markdown file for the rag-kit ingestion pipeline.\n"
        "It contains both English and Chinese content.\n\n"
        "## 安装说明\n\n"
        "使用以下命令安装：\n\n"
        "```bash\n"
        "pip install rag-kit\n"
        "```\n\n"
        "## Usage\n\n"
        "Run the following to start the watcher:\n\n"
        "```python\n"
        "from rag_kit import ingest\n"
        "chunks = ingest.ingest_folder('./docs')\n"
        "```\n"
    )
    path.write_text(content, encoding="utf-8")


def _setup_fixtures(tmp_dir: Path) -> dict[str, Path]:
    """Create all test fixtures in tmp_dir. Returns dict of name -> path."""
    fixtures = {}
    fixtures["text_pdf"] = tmp_dir / "text_guide.pdf"
    fixtures["scanned_pdf"] = tmp_dir / "scanned_doc.pdf"
    fixtures["docx"] = tmp_dir / "report.docx"
    fixtures["txt"] = tmp_dir / "notes.txt"
    fixtures["md"] = tmp_dir / "readme.md"

    _create_text_pdf(fixtures["text_pdf"])
    _create_scanned_pdf(fixtures["scanned_pdf"])
    _create_docx(fixtures["docx"])
    _create_txt(fixtures["txt"])
    _create_md(fixtures["md"])

    return fixtures


def main() -> int:
    print("=" * 60)
    print("rag-kit ingestion pipeline test")
    print("=" * 60)

    failures: list[str] = []
    tmp_dir = Path(tempfile.mkdtemp(prefix="rag_ingest_test_"))

    try:
        # --- 1. Create test fixtures ---
        print("\n[1] Creating test fixtures...")
        fixtures = _setup_fixtures(tmp_dir)
        for name, path in fixtures.items():
            print(f"    {name}: {path.name} ({path.stat().st_size} bytes)")

        # --- 2. Test individual file ingestion ---
        print("\n[2] Testing individual file ingestion...")

        # 2a. Text-based PDF
        print("\n  [2a] Text PDF (text_guide.pdf)...")
        chunks = ingest_file(fixtures["text_pdf"], chunk_size=256, chunk_overlap=32)
        if not chunks:
            failures.append("Text PDF: no chunks produced")
        else:
            print(f"      {len(chunks)} chunks extracted")
            # Verify text content is present
            all_text = " ".join(c["text"] for c in chunks)
            if "Hardware" not in all_text and "hardware" not in all_text:
                failures.append(f"Text PDF: English content missing")
            if "系统" not in all_text:
                failures.append("Text PDF: Chinese content missing")
            # Verify chunk structure
            for c in chunks:
                _validate_chunk(c, failures, "text_pdf")
            print(f"      ✓ Text PDF chunks valid (en+zh content present)")

        # 2b. Scanned PDF (OCR may or may not be available)
        print("\n  [2b] Scanned PDF (scanned_doc.pdf)...")
        from rag_kit.ingest import is_ocr_available
        ocr_avail = is_ocr_available()
        chunks = ingest_file(
            fixtures["scanned_pdf"], chunk_size=256, chunk_overlap=32,
            use_ocr=ocr_avail,
        )
        if ocr_avail:
            if not chunks:
                failures.append("Scanned PDF: OCR available but no chunks produced")
            else:
                print(f"      {len(chunks)} chunks extracted (via OCR)")
                for c in chunks:
                    _validate_chunk(c, failures, "scanned_pdf")
                print(f"      ✓ Scanned PDF chunks valid (OCR)")
        else:
            print(f"      [SKIP] OCR not installed — scanned PDF returns no chunks")
            if chunks:
                failures.append("Scanned PDF: got chunks without OCR (unexpected)")

        # 2c. DOCX
        print("\n  [2c] DOCX (report.docx)...")
        chunks = ingest_file(fixtures["docx"], chunk_size=256, chunk_overlap=32)
        if not chunks:
            failures.append("DOCX: no chunks produced")
        else:
            print(f"      {len(chunks)} chunks extracted")
            all_text = " ".join(c["text"] for c in chunks)
            if "DDR4" not in all_text:
                failures.append("DOCX: table content (DDR4) missing")
            if "Conclusion" not in all_text:
                failures.append("DOCX: paragraph content missing")
            for c in chunks:
                _validate_chunk(c, failures, "docx")
            print(f"      ✓ DOCX chunks valid (paragraphs + table)")

        # 2d. TXT
        print("\n  [2d] TXT (notes.txt)...")
        chunks = ingest_file(fixtures["txt"], chunk_size=256, chunk_overlap=32)
        if not chunks:
            failures.append("TXT: no chunks produced")
        else:
            print(f"      {len(chunks)} chunks extracted")
            all_text = " ".join(c["text"] for c in chunks)
            if "故障" not in all_text:
                failures.append("TXT: Chinese content missing")
            if "Device" not in all_text:
                failures.append("TXT: English content missing")
            for c in chunks:
                _validate_chunk(c, failures, "txt")
            print(f"      ✓ TXT chunks valid (zh+en content)")

        # 2e. MD
        print("\n  [2e] MD (readme.md)...")
        chunks = ingest_file(fixtures["md"], chunk_size=256, chunk_overlap=32)
        if not chunks:
            failures.append("MD: no chunks produced")
        else:
            print(f"      {len(chunks)} chunks extracted")
            all_text = " ".join(c["text"] for c in chunks)
            if "pip install" not in all_text:
                failures.append("MD: code block content missing")
            if "安装" not in all_text:
                failures.append("MD: Chinese content missing")
            for c in chunks:
                _validate_chunk(c, failures, "md")
            print(f"      ✓ MD chunks valid (headings + code + zh)")

        # --- 3. Test folder ingestion ---
        print("\n[3] Testing folder ingestion...")
        all_chunks = ingest_folder(tmp_dir, chunk_size=256, chunk_overlap=32)
        print(f"    Total chunks from folder: {len(all_chunks)}")

        source_files = set(c["source_file"] for c in all_chunks)
        expected_files = {"text_guide.pdf", "report.docx", "notes.txt", "readme.md"}
        if not ocr_avail:
            expected_files.discard("scanned_doc.pdf")
        # Scanned PDF may or may not produce chunks.
        if "scanned_doc.pdf" not in source_files and ocr_avail:
            failures.append("Folder ingestion: scanned_doc.pdf missing from results")

        for f in expected_files:
            if f not in source_files:
                failures.append(f"Folder ingestion: {f} missing from results")

        if len(expected_files - source_files) == 0:
            print(f"    ✓ All expected files present in results")

        # --- 4. Test language detection ---
        print("\n[4] Testing language detection in chunks...")
        lang_counts: dict[str, int] = {}
        for c in all_chunks:
            lang = c.get("language_detected", "unknown")
            lang_counts[lang] = lang_counts.get(lang, 0) + 1
        print(f"    Language distribution: {lang_counts}")

        if "en" not in lang_counts and "zh-en" not in lang_counts:
            failures.append("Language detection: no English chunks detected")
        if "zh" not in lang_counts and "zh-en" not in lang_counts:
            failures.append("Language detection: no Chinese chunks detected")
        if lang_counts:
            print(f"    ✓ Language detection working (zh/en/zh-en)")

        # --- 5. Test encoding integrity ---
        print("\n[5] Testing UTF-8 encoding integrity...")
        encoding_ok = True
        for c in all_chunks:
            text = c["text"]
            # Verify text is valid UTF-8 (re-encode + decode round-trip).
            try:
                text.encode("utf-8").decode("utf-8")
            except UnicodeError:
                failures.append(f"Encoding error in chunk {c['id']}")
                encoding_ok = False
        if encoding_ok:
            print(f"    ✓ All {len(all_chunks)} chunks pass UTF-8 round-trip")

        # --- 6. Test unsupported file handling ---
        print("\n[6] Testing unsupported file handling...")
        unsupported = tmp_dir / "data.csv"
        unsupported.write_text("a,b,c\n1,2,3\n", encoding="utf-8")
        try:
            ingest_file(unsupported)
            failures.append("Unsupported file: should have raised ValueError")
        except ValueError:
            print(f"    ✓ Unsupported file (.csv) correctly rejected")

        # --- 7. Test file-not-found handling ---
        print("\n[7] Testing file-not-found handling...")
        try:
            ingest_file(tmp_dir / "nonexistent.pdf")
            failures.append("Missing file: should have raised FileNotFoundError")
        except FileNotFoundError:
            print(f"    ✓ Missing file correctly raises FileNotFoundError")

        # --- 8. Test empty file handling ---
        print("\n[8] Testing empty file handling...")
        empty_txt = tmp_dir / "empty.txt"
        empty_txt.write_text("", encoding="utf-8")
        chunks = ingest_file(empty_txt)
        if chunks:
            failures.append("Empty file: should return empty list")
        else:
            print(f"    ✓ Empty file returns empty list")

    except Exception as exc:
        import traceback
        traceback.print_exc()
        failures.append(f"Exception: {exc}")

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    # --- Summary ---
    print("\n" + "=" * 60)
    if failures:
        print(f"FAILED — {len(failures)} issue(s):")
        for f in failures:
            print(f"  ✗ {f}")
        return 1
    else:
        print("ALL CHECKS PASSED ✓")
        print("  - Text PDF: Chinese + English content extracted")
        print("  - Scanned PDF: OCR fallback tested")
        print("  - DOCX: Paragraphs + table extracted")
        print("  - TXT: Chinese + English content extracted")
        print("  - MD: Headings + code blocks extracted")
        print("  - Folder ingestion: all files processed")
        print("  - Language detection: zh/en/zh-en working")
        print("  - UTF-8 encoding: no errors")
        print("  - Error handling: unsupported/missing/empty files")
        return 0


def _validate_chunk(chunk: dict, failures: list[str], label: str) -> None:
    """Validate that a chunk dict has all required keys with correct types."""
    required_keys = {"id", "source_file", "page", "section", "text", "language_detected"}
    missing = required_keys - set(chunk.keys())
    if missing:
        failures.append(f"{label}: chunk missing keys: {missing}")
        return

    if not chunk["id"]:
        failures.append(f"{label}: chunk has empty id")
    if not chunk["source_file"]:
        failures.append(f"{label}: chunk has empty source_file")
    if not isinstance(chunk["page"], int):
        failures.append(f"{label}: page is not int ({type(chunk['page'])})")
    if not isinstance(chunk["text"], str) or not chunk["text"].strip():
        failures.append(f"{label}: chunk has empty text")
    if chunk["language_detected"] not in ("zh", "en", "zh-en", "unknown"):
        failures.append(
            f"{label}: invalid language_detected '{chunk['language_detected']}'"
        )


if __name__ == "__main__":
    raise SystemExit(main())
